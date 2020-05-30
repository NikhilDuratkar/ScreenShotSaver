from docx import Document
from docx.shared import Inches
import keyboard
import time
import pyautogui
import os
import datetime

image_folderpath = "C:/Users/Mr.N/Documents/Lightshot/"
file_name = "ScreenShots_" + str(datetime.datetime.date(datetime.datetime.now())) + ".docx"

if file_name not in os.listdir(os.getcwd()):
    document_1 = Document()
    document_1.add_heading("All Today's Screenshots are appended in this docx.")
    document_1.save(file_name)
    print("Use 'Alt + q' to take Screenshot")
    print("Created new File - " + file_name + ". Image will be appended in it.\n")
else:
    print("Use 'Alt + q' to take Screenshot")
    print("Opened " + file_name + " File. Image will be appended in it.\n")

document = Document(file_name)


def add_image():
    screen_shot = 'screenshot.png'
    document.add_picture(screen_shot, width=Inches(7))
    print("ScreenShot added in", file_name)
    document.save(file_name)
    print(file_name + " Saved\n")


def take_image():
    image = pyautogui.screenshot()
    image.save("screenshot.png")


while True:
    try:
        if keyboard.read_key() == "alt" and keyboard.read_key() == "q":
            print("Alt+q key press detected")
            time.sleep(0.5)
            take_image()
            add_image()

    except Exception as e:
        print(e)
